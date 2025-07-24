import os
import tempfile
import streamlit as st
from pypdf import PdfReader
import whisper
import difflib
from moviepy.editor import VideoFileClip
from docx import Document

# -------------------------
# Transcribe video with Whisper
# -------------------------
def transcribe_video(video_path):
    try:
        audio_path = convert_video_to_audio(video_path)
        model = whisper.load_model("base")
        result = model.transcribe(audio_path)
        return result["text"]
    except Exception as e:
        return f"An unexpected error occurred during transcription: {e}"

# -------------------------
# Convert video to audio
# -------------------------
def convert_video_to_audio(video_path):
    audio_path = os.path.splitext(video_path)[0] + ".mp3"
    clip = VideoFileClip(video_path)
    clip.audio.write_audiofile(audio_path)
    return audio_path

# -------------------------
# Extract text + tables from PDF report
# -------------------------
def parse_pdf_report(file_path):
    try:
        reader = PdfReader(file_path)
        full_text = ""

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            full_text += f"\n--- Page {page_num + 1} ---\n"
            full_text += text if text else ""

            # Attempt to find table-like rows
            lines = text.splitlines() if text else []
            table_lines = [line for line in lines if '|' in line or '\t' in line or is_probable_table_row(line)]

            if table_lines:
                full_text += "\n\n[Extracted Table Data]\n"
                for line in table_lines:
                    full_text += line + "\n"

        return full_text.strip()

    except Exception as e:
        return f"[Error while parsing PDF report: {e}]"

def is_probable_table_row(line):
    tokens = line.split()
    return len(tokens) >= 3 and all(any(c.isdigit() for c in token) or token.isalpha() for token in tokens)

# -------------------------
# Compare video transcript with report
# -------------------------
def compare_texts(transcript, report):
    diff = difflib.ndiff(report.split(), transcript.split())
    comparison = "\n".join(diff)
    return comparison

# -------------------------
# Export comparison to Word
# -------------------------
def export_to_word(transcript, report, comparison):
    doc = Document()
    doc.add_heading("DUI Case Analysis Report", 0)

    doc.add_heading("Transcription from Video", level=1)
    doc.add_paragraph(transcript)

    doc.add_heading("Police Report Text (with Tables)", level=1)
    doc.add_paragraph(report)

    doc.add_heading("Comparison", level=1)
    doc.add_paragraph(comparison)

    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_docx.name)
    return temp_docx.name

# -------------------------
# Streamlit UI
# -------------------------
st.set_page_config(page_title="DUI Case Analyzer", layout="wide")
st.title("🚓 DUI Case Analyzer (Video + Report Comparator)")

video_file = st.file_uploader("Upload Bodycam Video", type=["mp4", "mov", "avi"])
report_file = st.file_uploader("Upload Police Report (PDF)", type=["pdf"])

if st.button("Analyze") and video_file and report_file:
    with st.spinner("Analyzing..."):
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp_vid:
                tmp_vid.write(video_file.read())
                tmp_video_path = tmp_vid.name

            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                tmp_pdf.write(report_file.read())
                tmp_pdf_path = tmp_pdf.name

            transcript = transcribe_video(tmp_video_path)
            report_text = parse_pdf_report(tmp_pdf_path)
            comparison = compare_texts(transcript, report_text)
            docx_path = export_to_word(transcript, report_text, comparison)

            st.subheader("Transcript")
            st.text_area("Video Transcript", transcript, height=200)

            st.subheader("Police Report")
            st.text_area("Report Content + Tables", report_text, height=200)

            st.subheader("Comparison")
            st.text_area("Comparison Result", comparison, height=300)

            with open(docx_path, "rb") as f:
                st.download_button(
                    label="📄 Download Full Report (Word)",
                    data=f,
                    file_name="dui_case_analysis.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"Something went wrong during analysis: {e}")
